"""
DOCX Report Exporter
====================

Generates court-ready DOCX forensic reports using python-docx.
"""

from __future__ import annotations

import textwrap
from datetime import UTC, datetime
from typing import Any

from core.structured_logging import get_logger

logger = get_logger(__name__)


def _normalize_report_for_template(report_dict: dict[str, Any]) -> dict[str, Any]:
    """Map DTO field names to template-friendly view-model keys."""
    return {
        "verdict": report_dict.get("overall_verdict", "INCONCLUSIVE"),
        "confidence": report_dict.get("overall_confidence", report_dict.get("confidence", 0.0)),
        "case_id": report_dict.get("case_id", "UNKNOWN"),
        "narrative": report_dict.get("executive_summary", report_dict.get("narrative", report_dict.get("summary", "No narrative generated."))),
        "findings": report_dict.get("per_agent_findings", report_dict.get("findings", {})),
        "report_hash": report_dict.get("report_hash", "N/A"),
        "manipulation_probability": report_dict.get("manipulation_probability", 0.0),
        "session_id": report_dict.get("session_id", ""),
        "signed_utc": report_dict.get("signed_utc", datetime.now(UTC).isoformat()),
        "cryptographic_signature": report_dict.get("cryptographic_signature", ""),
    }


async def export_report_docx(report_dict: dict[str, Any], session_id: str) -> bytes | None:
    """Export a ForensicReport as a DOCX file (bytes)."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        logger.warning("python-docx is not installed")
        return None

    tmpl = _normalize_report_for_template(report_dict)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    # ── Cover Page ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FORENSIC COUNCIL")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Multi-Agent Forensic Evidence Analysis Report")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0xAA)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Case ID: {tmpl['case_id']}\n").font.size = Pt(9)
    meta.add_run(f"Session: {session_id}\n").font.size = Pt(9)
    meta.add_run(f"Generated: {tmpl['signed_utc']}").font.size = Pt(9)

    doc.add_paragraph()
    verdict_p = doc.add_paragraph()
    verdict_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = verdict_p.add_run(f"VERDICT: {tmpl['verdict']}")
    run.bold = True
    run.font.size = Pt(16)
    prob = float(tmpl["manipulation_probability"])
    doc.add_paragraph().add_run(f"Manipulation Probability: {prob:.1%}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── Executive Summary ───────────────────────────────────────────────────
    doc.add_heading("Investigation Narrative", level=1)
    for line in textwrap.wrap(str(tmpl["narrative"]), width=90):
        doc.add_paragraph(line)

    doc.add_paragraph()

    # ── Agent Findings Table ────────────────────────────────────────────────
    doc.add_heading("Agent Findings", level=1)
    findings_raw = tmpl["findings"]
    all_findings: list[dict] = []
    if isinstance(findings_raw, dict):
        for agent_id, agent_findings in findings_raw.items():
            if isinstance(agent_findings, list):
                for f in agent_findings:
                    if isinstance(f, dict):
                        all_findings.append({**f, "_agent_id": agent_id})
    elif isinstance(findings_raw, list):
        all_findings = findings_raw

    if all_findings:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        headers = ["Agent", "Finding Type", "Verdict", "Confidence", "Severity"]
        for i, h in enumerate(headers):
            hdr[i].text = h
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(8)

        for f in all_findings[:20]:
            row = table.add_row().cells
            row[0].text = str(f.get("_agent_id", f.get("agent_id", "")))[:30]
            row[1].text = str(f.get("finding_type", ""))[:40]
            row[2].text = str(f.get("evidence_verdict", "INCONCLUSIVE"))[:20]
            conf = f.get("confidence_raw") or f.get("raw_confidence_score")
            row[3].text = f"{float(conf):.1%}" if conf is not None else "N/A"
            row[4].text = str(f.get("severity_tier", "INFO"))[:15]
            for cell in row:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        if len(all_findings) > 20:
            doc.add_paragraph(f"... and {len(all_findings) - 20} more findings.")

    doc.add_paragraph()

    # ── Chain of Custody / Integrity ────────────────────────────────────────
    doc.add_heading("Integrity & Chain of Custody", level=1)
    doc.add_heading("Report Hash (SHA-256)", level=2)
    p = doc.add_paragraph(tmpl["report_hash"])
    p.style = doc.styles["Normal"]
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(8)

    doc.add_heading("Custody Metadata", level=2)
    custody_table = doc.add_table(rows=6, cols=2)
    custody_table.style = "Light Grid Accent 1"
    custody_data = [
        ("Session ID", session_id),
        ("Case ID", tmpl["case_id"]),
        ("Report Generated", tmpl["signed_utc"]),
        ("Verdict", tmpl["verdict"]),
        ("Manipulation Probability", f"{prob:.1%}"),
        ("Cryptographic Signature", tmpl["cryptographic_signature"][:40] or "N/A"),
    ]
    for i, (key, val) in enumerate(custody_data):
        custody_table.rows[i].cells[0].text = key
        custody_table.rows[i].cells[1].text = str(val)
        for cell in custody_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    # ── Footer ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("FORENSIC COUNCIL — CONFIDENTIAL FORENSIC REPORT")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    footer_p2 = doc.add_paragraph()
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = footer_p2.add_run("This report is generated by an automated multi-agent forensic system. All findings should be reviewed by a qualified forensic analyst before use in legal proceedings.")
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
